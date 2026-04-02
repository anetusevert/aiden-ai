"""DOCX export service for legal workflow results.

This module provides enterprise-grade DOCX generation for Contract Review
and Clause Redlines workflows with full traceability and audit support.

Features:
- Professional Word document formatting with no external dependencies
- Full traceability footer with LLM provider, model, and prompt hash
- Structured citations with evidence appendix
- Deterministic output (same input = same document structure)
"""

import io
from datetime import datetime, timezone
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Pt, RGBColor

from src.config import settings
from src.schemas.clause_redlines import (
    ClauseRedlineItem,
    ClauseRedlinesResponse,
)
from src.schemas.contract_review import (
    ContractReviewResponse,
    EvidenceChunkRef,
    Finding,
)
from src.schemas.export import (
    CONFIDENTIALITY_NOTICE,
    LEGAL_DISCLAIMER,
    DocumentMetadata,
    WorkflowType,
)
from src.schemas.workflow_status import WorkflowResultStatus


class ExportService:
    """Service for generating DOCX exports of workflow results."""

    def __init__(self) -> None:
        """Initialize the export service."""
        self._environment = settings.environment

    def _create_document(self) -> Document:
        """Create a new Word document with default styling."""
        doc = Document()

        # Set default font for the document
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)

        return doc

    def _add_cover_page(
        self,
        doc: Document,
        title: str,
        metadata: DocumentMetadata,
        generated_at: datetime,
    ) -> None:
        """Add a cover page to the document.

        Args:
            doc: The Word document
            title: Memo title (e.g., "Contract Review Memo")
            metadata: Document and workspace metadata
            generated_at: When the export was generated
        """
        # Title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(24)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # Spacer

        # Document Title
        doc_title_para = doc.add_paragraph()
        doc_title_run = doc_title_para.add_run(f"Document: {metadata.document_title}")
        doc_title_run.bold = True
        doc_title_run.font.size = Pt(14)
        doc_title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Version
        version_para = doc.add_paragraph()
        version_run = version_para.add_run(f"Version: {metadata.version_number}")
        version_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # Spacer

        # Metadata table
        table = doc.add_table(rows=6, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Set column widths
        for row in table.rows:
            row.cells[0].width = Inches(2)
            row.cells[1].width = Inches(4)

        metadata_items = [
            ("Workspace", metadata.workspace_name),
            ("Organization", metadata.tenant_name),
            ("Jurisdiction", metadata.jurisdiction),
            ("Document ID", metadata.document_id[:8] + "..."),
            ("Version ID", metadata.version_id[:8] + "..."),
            ("Generated (UTC)", generated_at.strftime("%Y-%m-%d %H:%M:%S")),
        ]

        for i, (label, value) in enumerate(metadata_items):
            label_cell = table.rows[i].cells[0]
            value_cell = table.rows[i].cells[1]
            label_cell.paragraphs[0].runs[0].bold = True if label_cell.paragraphs[0].runs else None
            label_para = label_cell.paragraphs[0]
            label_para.clear()
            label_run = label_para.add_run(label + ":")
            label_run.bold = True
            value_cell.text = value

        doc.add_paragraph()  # Spacer
        doc.add_paragraph()  # Spacer

        # Confidentiality notice
        conf_para = doc.add_paragraph()
        conf_run = conf_para.add_run(CONFIDENTIALITY_NOTICE)
        conf_run.font.size = Pt(9)
        conf_run.italic = True
        conf_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Page break after cover
        doc.add_page_break()

    def _add_executive_summary(
        self,
        doc: Document,
        summary: str,
        insufficient_sources: bool,
    ) -> None:
        """Add the executive summary section.

        Args:
            doc: The Word document
            summary: Summary text from the workflow
            insufficient_sources: Whether the workflow had insufficient sources
        """
        heading = doc.add_heading("Executive Summary", level=1)

        # Insufficient sources disclaimer
        if insufficient_sources:
            disclaimer_para = doc.add_paragraph()
            disclaimer_run = disclaimer_para.add_run(
                "NOTICE: This analysis was generated with insufficient source material. "
                "The findings below may be incomplete or less reliable than a fully-sourced review. "
                "Please ensure all relevant documents are uploaded and indexed for comprehensive analysis."
            )
            disclaimer_run.bold = True
            disclaimer_run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)  # Dark red
            doc.add_paragraph()  # Spacer

        # Summary text
        doc.add_paragraph(summary)
        doc.add_paragraph()  # Spacer

    def _add_contract_review_findings(
        self,
        doc: Document,
        findings: list[Finding],
    ) -> None:
        """Add the findings section for contract review.

        Args:
            doc: The Word document
            findings: List of findings from the contract review
        """
        doc.add_heading("Findings", level=1)

        if not findings:
            doc.add_paragraph("No findings were identified in this review.")
            return

        for i, finding in enumerate(findings, 1):
            # Finding heading
            finding_heading = doc.add_heading(
                f"Finding {i}: {finding.title}", level=2
            )

            # Metadata table
            table = doc.add_table(rows=2, cols=4)
            table.style = "Table Grid"

            # Row 1: Severity and Category labels
            table.rows[0].cells[0].text = "Severity"
            table.rows[0].cells[1].text = finding.severity.upper()
            table.rows[0].cells[2].text = "Category"
            table.rows[0].cells[3].text = finding.category.replace("_", " ").title()

            # Apply severity color
            severity_colors = {
                "critical": "DC2626",
                "high": "EA580C",
                "medium": "F59E0B",
                "low": "3B82F6",
            }
            color = severity_colors.get(finding.severity, "6B7280")
            self._set_cell_background(table.rows[0].cells[1], color)

            # Make header cells bold
            for cell in table.rows[0].cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True

            doc.add_paragraph()  # Spacer

            # Issue
            issue_heading = doc.add_paragraph()
            issue_heading_run = issue_heading.add_run("Issue:")
            issue_heading_run.bold = True
            doc.add_paragraph(finding.issue)

            # Recommendation
            rec_heading = doc.add_paragraph()
            rec_heading_run = rec_heading.add_run("Recommendation:")
            rec_heading_run.bold = True
            doc.add_paragraph(finding.recommendation)

            # Citations
            if finding.citations:
                citations_para = doc.add_paragraph()
                citations_run = citations_para.add_run(
                    f"Citations: [{', '.join(str(c) for c in finding.citations)}]"
                )
                citations_run.italic = True

            doc.add_paragraph()  # Spacer between findings

    def _add_clause_redlines_items(
        self,
        doc: Document,
        items: list[ClauseRedlineItem],
    ) -> None:
        """Add the clause redlines section.

        Args:
            doc: The Word document
            items: List of clause redline items
        """
        doc.add_heading("Clause Analysis", level=1)

        if not items:
            doc.add_paragraph("No clause items were identified in this analysis.")
            return

        for item in items:
            # Clause type heading
            clause_label = item.clause_type.replace("_", " ").title()
            doc.add_heading(clause_label, level=2)

            # Status and metadata table
            table = doc.add_table(rows=2, cols=4)
            table.style = "Table Grid"

            # Row 1: Status and Severity
            table.rows[0].cells[0].text = "Status"
            table.rows[0].cells[1].text = item.status.replace("_", " ").upper()
            table.rows[0].cells[2].text = "Severity"
            table.rows[0].cells[3].text = item.severity.upper()

            # Row 2: Confidence
            table.rows[1].cells[0].text = "Confidence"
            table.rows[1].cells[1].text = f"{int(item.confidence * 100)}%"
            table.rows[1].cells[2].text = "Level"
            table.rows[1].cells[3].text = item.confidence_level.upper()

            # Apply status color
            status_colors = {
                "found": "22C55E",
                "missing": "EF4444",
                "insufficient_evidence": "F59E0B",
            }
            color = status_colors.get(item.status, "6B7280")
            self._set_cell_background(table.rows[0].cells[1], color)

            # Make header cells bold
            for row in table.rows:
                for i in [0, 2]:  # Label columns
                    for para in row.cells[i].paragraphs:
                        for run in para.runs:
                            run.bold = True

            doc.add_paragraph()  # Spacer

            # Issue (if present)
            if item.issue:
                issue_heading = doc.add_paragraph()
                issue_heading_run = issue_heading.add_run("Issue:")
                issue_heading_run.bold = True
                doc.add_paragraph(item.issue)

            # Rationale (if present)
            if item.rationale:
                rationale_heading = doc.add_paragraph()
                rationale_heading_run = rationale_heading.add_run("Rationale:")
                rationale_heading_run.bold = True
                doc.add_paragraph(item.rationale)

            # Suggested redline (if present)
            if item.suggested_redline:
                redline_heading = doc.add_paragraph()
                redline_heading_run = redline_heading.add_run("Recommended Text:")
                redline_heading_run.bold = True
                redline_heading_run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)  # Blue

                redline_para = doc.add_paragraph()
                redline_para.paragraph_format.left_indent = Inches(0.25)
                redline_para.add_run(item.suggested_redline)

            # Citations
            if item.citations:
                citations_para = doc.add_paragraph()
                citations_run = citations_para.add_run(
                    f"Citations: [{', '.join(str(c) for c in item.citations)}]"
                )
                citations_run.italic = True

            doc.add_paragraph()  # Spacer between items

    def _add_citations_section(
        self,
        doc: Document,
        evidence_chunks: list[EvidenceChunkRef],
        document_title: str,
        version_number: int,
    ) -> None:
        """Add the citations section with numbered list.

        Supports both workspace documents and global legal sources with explicit provenance.

        Args:
            doc: The Word document
            evidence_chunks: List of evidence chunk references
            document_title: Title of the source document (fallback for workspace)
            version_number: Version number of the source document
        """
        doc.add_heading("Citations", level=1)

        if not evidence_chunks:
            doc.add_paragraph("No citations available.")
            return

        for i, chunk in enumerate(evidence_chunks, 1):
            # Citation entry
            citation_para = doc.add_paragraph()
            citation_run = citation_para.add_run(f"[{i}] ")
            citation_run.bold = True

            # Determine source info based on source_type
            source_type = getattr(chunk, 'source_type', 'workspace_document')
            source_label = getattr(chunk, 'source_label', '')

            if source_type == "global_legal":
                # Global legal source
                source_text = f"Source: {source_label or 'Global Law'}"
                source_badge = " [GLOBAL LAW]"
                jurisdiction = getattr(chunk, 'jurisdiction', None)
                official_url = getattr(chunk, 'official_source_url', None)

                citation_para.add_run(source_text)
                badge_run = citation_para.add_run(source_badge)
                badge_run.bold = True
                badge_run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)  # Blue

                if jurisdiction:
                    citation_para.add_run(f" | Jurisdiction: {jurisdiction}")
                if official_url:
                    url_para = doc.add_paragraph()
                    url_para.paragraph_format.left_indent = Inches(0.25)
                    url_run = url_para.add_run(f"Official source: {official_url}")
                    url_run.font.size = Pt(9)
                    url_run.italic = True
            else:
                # Workspace document source
                source_text = f"Source: {source_label or document_title}, Version {version_number}"
                source_badge = " [WORKSPACE DOC]"

                citation_para.add_run(source_text)
                badge_run = citation_para.add_run(source_badge)
                badge_run.font.size = Pt(9)
                badge_run.italic = True

            # Chunk info
            chunk_info = f" (Chunk: {chunk.chunk_id[:8]}..., Chars {chunk.char_start}-{chunk.char_end})"
            chunk_run = citation_para.add_run(chunk_info)
            chunk_run.font.size = Pt(9)
            chunk_run.italic = True

            # Snippet excerpt (truncated)
            snippet_excerpt = chunk.snippet[:200] + "..." if len(chunk.snippet) > 200 else chunk.snippet
            excerpt_para = doc.add_paragraph()
            excerpt_para.paragraph_format.left_indent = Inches(0.25)
            excerpt_run = excerpt_para.add_run(f'"{snippet_excerpt}"')
            excerpt_run.italic = True
            excerpt_run.font.size = Pt(10)

        doc.add_paragraph()  # Spacer

    def _add_evidence_appendix(
        self,
        doc: Document,
        evidence_chunks: list[EvidenceChunkRef],
        document_title: str,
        version_number: int,
    ) -> None:
        """Add the evidence appendix with full snippets.

        Supports both workspace documents and global legal sources with explicit provenance.

        Args:
            doc: The Word document
            evidence_chunks: List of evidence chunk references
            document_title: Title of the source document (fallback for workspace)
            version_number: Version number of the source document
        """
        doc.add_page_break()
        doc.add_heading("Evidence Appendix", level=1)

        if not evidence_chunks:
            doc.add_paragraph("No evidence available.")
            return

        for i, chunk in enumerate(evidence_chunks, 1):
            # Determine source info based on source_type
            source_type = getattr(chunk, 'source_type', 'workspace_document')
            source_label = getattr(chunk, 'source_label', '')
            jurisdiction = getattr(chunk, 'jurisdiction', None)
            official_url = getattr(chunk, 'official_source_url', None)

            # Evidence entry heading with source badge
            if source_type == "global_legal":
                doc.add_heading(f"Evidence [{i}] — GLOBAL LAW", level=2)
            else:
                doc.add_heading(f"Evidence [{i}] — WORKSPACE DOC", level=2)

            # Metadata table - adjust rows based on source type
            num_rows = 4 if source_type == "global_legal" else 3
            table = doc.add_table(rows=num_rows, cols=2)
            table.style = "Table Grid"

            table.rows[0].cells[0].text = "Chunk ID"
            table.rows[0].cells[1].text = chunk.chunk_id
            table.rows[1].cells[0].text = "Character Range"
            table.rows[1].cells[1].text = f"{chunk.char_start} - {chunk.char_end}"

            if source_type == "global_legal":
                table.rows[2].cells[0].text = "Source"
                table.rows[2].cells[1].text = source_label or "Global Legal Corpus"
                table.rows[3].cells[0].text = "Jurisdiction"
                table.rows[3].cells[1].text = jurisdiction or "N/A"
            else:
                table.rows[2].cells[0].text = "Source"
                table.rows[2].cells[1].text = f"{source_label or document_title} → Version {version_number}"

            # Make label cells bold
            for row in table.rows:
                for para in row.cells[0].paragraphs:
                    for run in para.runs:
                        run.bold = True

            doc.add_paragraph()  # Spacer

            # Official source URL for global legal
            if source_type == "global_legal" and official_url:
                url_heading = doc.add_paragraph()
                url_heading_run = url_heading.add_run("Official Source URL:")
                url_heading_run.bold = True
                url_para = doc.add_paragraph()
                url_para.paragraph_format.left_indent = Inches(0.25)
                url_run = url_para.add_run(official_url)
                url_run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)  # Blue

            # Full snippet
            snippet_heading = doc.add_paragraph()
            snippet_heading_run = snippet_heading.add_run("Full Text:")
            snippet_heading_run.bold = True

            snippet_para = doc.add_paragraph()
            snippet_para.paragraph_format.left_indent = Inches(0.25)
            snippet_para.add_run(chunk.snippet)

            doc.add_paragraph()  # Spacer between entries

    def _add_traceability_footer(
        self,
        doc: Document,
        workflow_name: str,
        status: WorkflowResultStatus,
        llm_provider: str | None,
        llm_model: str | None,
        prompt_hash: str | None,
    ) -> None:
        """Add the mandatory traceability footer.

        Args:
            doc: The Word document
            workflow_name: Name of the workflow (e.g., CONTRACT_REVIEW_V1)
            status: Workflow result status
            llm_provider: LLM provider used
            llm_model: LLM model used
            prompt_hash: SHA256 hash of the prompt
        """
        doc.add_page_break()
        doc.add_heading("Traceability Information", level=1)

        # Traceability table
        table = doc.add_table(rows=7, cols=2)
        table.style = "Table Grid"

        traceability_items = [
            ("Workflow", workflow_name),
            ("Result Status", status.value),
            ("LLM Provider", llm_provider or "N/A"),
            ("LLM Model", llm_model or "N/A"),
            ("Prompt Hash", prompt_hash[:16] + "..." if prompt_hash else "N/A"),
            ("Generated By", "Aiden.ai"),
            ("Environment", self._environment),
        ]

        for i, (label, value) in enumerate(traceability_items):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = value

            # Make label cells bold
            for para in table.rows[i].cells[0].paragraphs:
                for run in para.runs:
                    run.bold = True

        doc.add_paragraph()  # Spacer

        # Legal disclaimer
        disclaimer_para = doc.add_paragraph()
        disclaimer_run = disclaimer_para.add_run(LEGAL_DISCLAIMER)
        disclaimer_run.font.size = Pt(9)
        disclaimer_run.italic = True

    def _set_cell_background(self, cell: Any, color: str) -> None:
        """Set the background color of a table cell.

        Args:
            cell: The table cell
            color: Hex color code (without #)
        """
        shading_elm = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="{color}"/>'
        )
        cell._tc.get_or_add_tcPr().append(shading_elm)

    def _collect_all_evidence(
        self, findings: list[Finding]
    ) -> list[EvidenceChunkRef]:
        """Collect all unique evidence chunks from findings.

        Args:
            findings: List of findings

        Returns:
            De-duplicated list of evidence chunks
        """
        seen_chunk_ids: set[str] = set()
        evidence: list[EvidenceChunkRef] = []

        for finding in findings:
            for chunk in finding.evidence:
                if chunk.chunk_id not in seen_chunk_ids:
                    seen_chunk_ids.add(chunk.chunk_id)
                    evidence.append(chunk)

        return evidence

    def _collect_all_evidence_from_items(
        self, items: list[ClauseRedlineItem]
    ) -> list[EvidenceChunkRef]:
        """Collect all unique evidence chunks from clause items.

        Args:
            items: List of clause redline items

        Returns:
            De-duplicated list of evidence chunks
        """
        seen_chunk_ids: set[str] = set()
        evidence: list[EvidenceChunkRef] = []

        for item in items:
            for chunk in item.evidence:
                if chunk.chunk_id not in seen_chunk_ids:
                    seen_chunk_ids.add(chunk.chunk_id)
                    # Convert to EvidenceChunkRef (same structure)
                    evidence.append(
                        EvidenceChunkRef(
                            chunk_id=chunk.chunk_id,
                            snippet=chunk.snippet,
                            char_start=chunk.char_start,
                            char_end=chunk.char_end,
                        )
                    )

        return evidence

    def generate_contract_review_docx(
        self,
        result: ContractReviewResponse,
        metadata: DocumentMetadata,
    ) -> bytes:
        """Generate a DOCX export for contract review results.

        Args:
            result: The contract review workflow result
            metadata: Document and workspace metadata

        Returns:
            DOCX file as bytes
        """
        generated_at = datetime.now(timezone.utc)
        doc = self._create_document()

        # Cover page
        self._add_cover_page(
            doc,
            "Contract Review Memo",
            metadata,
            generated_at,
        )

        # Executive summary
        self._add_executive_summary(
            doc,
            result.summary,
            result.insufficient_sources,
        )

        # Findings
        self._add_contract_review_findings(doc, result.findings)

        # Collect all evidence
        all_evidence = self._collect_all_evidence(result.findings)

        # Citations section
        self._add_citations_section(
            doc,
            all_evidence,
            metadata.document_title,
            metadata.version_number,
        )

        # Evidence appendix
        self._add_evidence_appendix(
            doc,
            all_evidence,
            metadata.document_title,
            metadata.version_number,
        )

        # Traceability footer
        self._add_traceability_footer(
            doc,
            WorkflowType.CONTRACT_REVIEW,
            result.meta.status,
            result.meta.llm_provider or result.meta.provider,
            result.meta.llm_model or result.meta.model,
            result.meta.prompt_hash,
        )

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def generate_clause_redlines_docx(
        self,
        result: ClauseRedlinesResponse,
        metadata: DocumentMetadata,
    ) -> bytes:
        """Generate a DOCX export for clause redlines results.

        Args:
            result: The clause redlines workflow result
            metadata: Document and workspace metadata

        Returns:
            DOCX file as bytes
        """
        generated_at = datetime.now(timezone.utc)
        doc = self._create_document()

        # Cover page
        self._add_cover_page(
            doc,
            "Clause Redlines Memo",
            metadata,
            generated_at,
        )

        # Executive summary
        self._add_executive_summary(
            doc,
            result.summary,
            result.insufficient_sources,
        )

        # Clause items
        self._add_clause_redlines_items(doc, result.items)

        # Collect all evidence
        all_evidence = self._collect_all_evidence_from_items(result.items)

        # Citations section
        self._add_citations_section(
            doc,
            all_evidence,
            metadata.document_title,
            metadata.version_number,
        )

        # Evidence appendix
        self._add_evidence_appendix(
            doc,
            all_evidence,
            metadata.document_title,
            metadata.version_number,
        )

        # Traceability footer
        self._add_traceability_footer(
            doc,
            WorkflowType.CLAUSE_REDLINES,
            result.meta.status,
            result.meta.llm_provider or result.meta.provider,
            result.meta.llm_model or result.meta.model,
            result.meta.prompt_hash,
        )

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
