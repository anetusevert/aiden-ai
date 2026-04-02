"""Text extraction from PDFs, DOCX, HTML, and plain text files.

Supports:
- PDF: PyMuPDF (fitz) preferred, pdfminer.six as fallback
- DOCX: python-docx
- HTML: BeautifulSoup
- Plain text: Direct decoding

No OCR is performed. Arabic text is stored as-is.
"""

import io
import logging
from dataclasses import dataclass
from typing import Literal

logger = logging.getLogger(__name__)

ExtractionMethod = Literal["pymupdf", "pdfminer", "docx", "html", "plain", "unsupported"]


@dataclass
class ExtractionResult:
    """Result of text extraction."""

    text: str
    page_count: int | None
    method: ExtractionMethod


def extract_pdf_text(file_bytes: bytes) -> ExtractionResult:
    """Extract text from a PDF file.

    Attempts PyMuPDF first, falls back to pdfminer.six if unavailable.
    No OCR is performed.

    Args:
        file_bytes: PDF file content as bytes

    Returns:
        ExtractionResult with text, page count, and method used
    """
    # Try PyMuPDF first (preferred - faster and better Arabic support)
    try:
        return _extract_pdf_with_pymupdf(file_bytes)
    except ImportError:
        logger.info("PyMuPDF not available, trying pdfminer.six")
    except Exception as e:
        logger.warning(f"PyMuPDF extraction failed, trying pdfminer.six: {e}")

    # Fall back to pdfminer.six
    try:
        return _extract_pdf_with_pdfminer(file_bytes)
    except ImportError:
        logger.error("Neither PyMuPDF nor pdfminer.six is available")
        return ExtractionResult(text="", page_count=None, method="unsupported")
    except Exception as e:
        logger.error(f"pdfminer.six extraction failed: {e}")
        return ExtractionResult(text="", page_count=None, method="pdfminer")


def _extract_pdf_with_pymupdf(file_bytes: bytes) -> ExtractionResult:
    """Extract text using PyMuPDF (fitz).

    Args:
        file_bytes: PDF file content

    Returns:
        ExtractionResult with text and page count
    """
    import fitz  # PyMuPDF

    doc = fitz.open(stream=file_bytes, filetype="pdf")
    try:
        text_parts: list[str] = []
        page_count = len(doc)

        for page in doc:
            # Extract text preserving layout reasonably
            # Don't use any text shaping options to preserve Arabic as-is
            page_text = page.get_text("text")
            if page_text:
                text_parts.append(page_text)

        # Join pages with double newline
        text = "\n\n".join(text_parts)

        return ExtractionResult(
            text=text,
            page_count=page_count,
            method="pymupdf",
        )
    finally:
        doc.close()


def _extract_pdf_with_pdfminer(file_bytes: bytes) -> ExtractionResult:
    """Extract text using pdfminer.six.

    Args:
        file_bytes: PDF file content

    Returns:
        ExtractionResult with text and page count
    """
    from pdfminer.high_level import extract_pages, extract_text as pdfminer_extract_text
    from pdfminer.layout import LAParams

    # Use default layout parameters - don't apply any transformations
    laparams = LAParams()

    # Extract text
    text = pdfminer_extract_text(io.BytesIO(file_bytes), laparams=laparams)

    # Count pages
    page_count = 0
    try:
        for _ in extract_pages(io.BytesIO(file_bytes)):
            page_count += 1
    except Exception:
        page_count = None

    return ExtractionResult(
        text=text if text else "",
        page_count=page_count,
        method="pdfminer",
    )


def extract_docx_text(file_bytes: bytes) -> ExtractionResult:
    """Extract text from a DOCX file.

    Uses python-docx to extract paragraph text.
    Arabic text is preserved as-is.

    Args:
        file_bytes: DOCX file content as bytes

    Returns:
        ExtractionResult with text and method
    """
    try:
        from docx import Document
    except ImportError:
        logger.error("python-docx is not available")
        return ExtractionResult(text="", page_count=None, method="unsupported")

    try:
        doc = Document(io.BytesIO(file_bytes))

        # Extract text from paragraphs
        paragraphs: list[str] = []
        for para in doc.paragraphs:
            if para.text:
                paragraphs.append(para.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text and cell.text not in paragraphs:
                        paragraphs.append(cell.text)

        # Join with newlines (preserve paragraph structure)
        text = "\n\n".join(paragraphs)

        return ExtractionResult(
            text=text,
            page_count=None,  # DOCX doesn't have fixed pages
            method="docx",
        )
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        return ExtractionResult(text="", page_count=None, method="docx")


def extract_html_text(file_bytes: bytes) -> ExtractionResult:
    """Extract text from an HTML file.

    Uses BeautifulSoup to parse HTML and extract text content.
    Removes script and style elements.
    Arabic text is preserved as-is.

    Args:
        file_bytes: HTML file content as bytes

    Returns:
        ExtractionResult with text and method
    """
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        logger.error("beautifulsoup4 is not available")
        return ExtractionResult(text="", page_count=None, method="unsupported")

    try:
        # Decode HTML - try utf-8 first, then fallback to latin-1
        try:
            html_content = file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            html_content = file_bytes.decode("latin-1")

        # Parse HTML
        soup = BeautifulSoup(html_content, "html.parser")

        # Remove script and style elements
        for element in soup(["script", "style", "head", "meta", "link"]):
            element.decompose()

        # Get text, separating blocks with newlines
        text = soup.get_text(separator="\n", strip=True)

        # Clean up excessive whitespace while preserving paragraph structure
        lines = [line.strip() for line in text.split("\n")]
        cleaned_lines = []
        prev_empty = False
        for line in lines:
            if line:
                cleaned_lines.append(line)
                prev_empty = False
            elif not prev_empty:
                cleaned_lines.append("")
                prev_empty = True

        text = "\n".join(cleaned_lines).strip()

        return ExtractionResult(
            text=text,
            page_count=None,  # HTML doesn't have pages
            method="html",
        )
    except Exception as e:
        logger.error(f"HTML extraction failed: {e}")
        return ExtractionResult(text="", page_count=None, method="html")


def extract_plain_text(file_bytes: bytes) -> ExtractionResult:
    """Extract text from a plain text file.

    Decodes bytes to text, trying utf-8 first, then fallback encodings.
    Arabic text is preserved as-is.

    Args:
        file_bytes: Plain text file content as bytes

    Returns:
        ExtractionResult with text and method
    """
    try:
        # Try common encodings
        for encoding in ["utf-8", "utf-16", "latin-1", "cp1256"]:  # cp1256 is Arabic Windows
            try:
                text = file_bytes.decode(encoding)
                return ExtractionResult(
                    text=text.strip(),
                    page_count=None,
                    method="plain",
                )
            except UnicodeDecodeError:
                continue

        # Last resort - decode with replacement
        text = file_bytes.decode("utf-8", errors="replace")
        return ExtractionResult(
            text=text.strip(),
            page_count=None,
            method="plain",
        )
    except Exception as e:
        logger.error(f"Plain text extraction failed: {e}")
        return ExtractionResult(text="", page_count=None, method="plain")


def extract_text(file_bytes: bytes, content_type: str) -> ExtractionResult:
    """Extract text from a file based on its content type.

    Args:
        file_bytes: File content as bytes
        content_type: MIME type of the file

    Returns:
        ExtractionResult with text and metadata
    """
    content_type_lower = content_type.lower()

    # PDF
    if content_type_lower == "application/pdf":
        return extract_pdf_text(file_bytes)

    # DOCX
    if content_type_lower in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/docx",
    ):
        return extract_docx_text(file_bytes)

    # HTML
    if content_type_lower in ("text/html", "application/xhtml+xml"):
        return extract_html_text(file_bytes)

    # Plain text
    if content_type_lower in ("text/plain",):
        return extract_plain_text(file_bytes)

    # Unsupported format - return empty
    logger.info(f"Unsupported content type for extraction: {content_type}")
    return ExtractionResult(text="", page_count=None, method="unsupported")
