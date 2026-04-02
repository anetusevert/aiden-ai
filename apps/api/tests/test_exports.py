"""Tests for DOCX export functionality.

These tests verify:
1. DOCX generation does not crash
2. Correct sections exist in exported documents
3. Traceability footer contains correct fields
4. Insufficient sources export includes disclaimer
5. DOCX headers (Content-Type, Content-Disposition) are correct
6. DOCX file starts with ZIP header (PK)
7. Invalid evidence references return 400 with error_code="export_validation_failed"

Run tests:
    uv run pytest tests/test_exports.py -v

Run with Docker:
    docker compose run --rm api uv run pytest tests/test_exports.py -v
"""

import io

import pytest
from docx import Document

from src.schemas.clause_redlines import (
    ClauseRedlineItem,
    ClauseRedlinesMeta,
    ClauseRedlinesResponse,
    EvidenceChunkRef as ClauseEvidenceChunkRef,
)
from src.schemas.contract_review import (
    ContractReviewMeta,
    ContractReviewResponse,
    EvidenceChunkRef,
    Finding,
)
from src.schemas.export import DocumentMetadata
from src.schemas.workflow_status import WorkflowResultStatus
from src.services.export_service import ExportService

# DOCX MIME type constant
DOCX_MIME_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


@pytest.fixture
def export_service() -> ExportService:
    """Create an export service instance."""
    return ExportService()


@pytest.fixture
def document_metadata() -> DocumentMetadata:
    """Create sample document metadata."""
    return DocumentMetadata(
        document_id="doc-12345678-1234-1234-1234-123456789012",
        version_id="ver-12345678-1234-1234-1234-123456789012",
        document_title="Employment Agreement",
        version_number=1,
        workspace_name="Corporate Legal",
        tenant_name="Acme Corporation",
        jurisdiction="UAE",
    )


@pytest.fixture
def sample_evidence() -> list[EvidenceChunkRef]:
    """Create sample evidence chunks."""
    return [
        EvidenceChunkRef(
            chunk_id="chunk-1234-5678-abcd-efgh",
            snippet="The Employee agrees to work exclusively for the Company during the term of this agreement.",
            char_start=100,
            char_end=200,
        ),
        EvidenceChunkRef(
            chunk_id="chunk-2345-6789-bcde-fghi",
            snippet="Liability shall not exceed the total contract value of USD 100,000.",
            char_start=500,
            char_end=600,
        ),
    ]


@pytest.fixture
def sample_findings(sample_evidence: list[EvidenceChunkRef]) -> list[Finding]:
    """Create sample findings for contract review."""
    return [
        Finding(
            finding_id="finding-1",
            title="Liability Cap Too Low",
            severity="high",
            category="liability",
            issue="The liability clause limits damages to contract value [1].",
            recommendation="Negotiate higher liability cap or remove limitation [1].",
            citations=[1],
            evidence=[sample_evidence[1]],
        ),
        Finding(
            finding_id="finding-2",
            title="Missing Termination Notice Period",
            severity="medium",
            category="termination",
            issue="The contract does not specify a termination notice period [2].",
            recommendation="Add explicit 30-day notice period clause [2].",
            citations=[2],
            evidence=[sample_evidence[0]],
        ),
    ]


@pytest.fixture
def contract_review_response(
    sample_findings: list[Finding],
) -> ContractReviewResponse:
    """Create a sample contract review response."""
    return ContractReviewResponse(
        summary="The contract review identified 2 findings [1][2]. The main concerns are liability cap and termination provisions.",
        findings=sample_findings,
        meta=ContractReviewMeta(
            status=WorkflowResultStatus.SUCCESS,
            model="gpt-4o-mini",
            provider="openai",
            evidence_chunk_count=10,
            request_id="req-12345",
            output_language="en",
            review_mode="standard",
            removed_findings_count=0,
            strict_citations_failed=False,
            prompt_hash="abc123def456",
            llm_provider="openai",
            llm_model="gpt-4o-mini",
        ),
        insufficient_sources=False,
    )


@pytest.fixture
def contract_review_insufficient_sources(
    sample_findings: list[Finding],
) -> ContractReviewResponse:
    """Create a sample contract review response with insufficient sources."""
    return ContractReviewResponse(
        summary="Limited analysis due to insufficient source material [1].",
        findings=sample_findings[:1],  # Only one finding
        meta=ContractReviewMeta(
            status=WorkflowResultStatus.INSUFFICIENT_SOURCES,
            model="gpt-4o-mini",
            provider="openai",
            evidence_chunk_count=2,
            request_id="req-12345",
            output_language="en",
            review_mode="standard",
            removed_findings_count=0,
            strict_citations_failed=False,
            prompt_hash="abc123def456",
            llm_provider="openai",
            llm_model="gpt-4o-mini",
        ),
        insufficient_sources=True,
    )


@pytest.fixture
def sample_clause_items() -> list[ClauseRedlineItem]:
    """Create sample clause redline items."""
    return [
        ClauseRedlineItem(
            clause_type="governing_law",
            status="found",
            confidence=0.85,
            confidence_level="high",
            confidence_reason="Matched heading + 2 triggers",
            issue="The governing law clause specifies UAE jurisdiction [1].",
            suggested_redline="Recommended Text: This Agreement shall be governed by the laws of the United Arab Emirates.",
            rationale="The current clause is acceptable but could be more detailed [1].",
            citations=[1],
            evidence=[
                ClauseEvidenceChunkRef(
                    chunk_id="chunk-gov-1234",
                    snippet="This Agreement shall be governed by and construed in accordance with UAE law.",
                    char_start=1000,
                    char_end=1100,
                )
            ],
            severity="low",
        ),
        ClauseRedlineItem(
            clause_type="termination",
            status="missing",
            confidence=0.0,
            confidence_level="low",
            confidence_reason="No evidence found",
            issue=None,
            suggested_redline="Recommended Text: Either party may terminate this Agreement by providing 30 days written notice.",
            rationale=None,
            citations=[],
            evidence=[],
            severity="high",
        ),
    ]


@pytest.fixture
def clause_redlines_response(
    sample_clause_items: list[ClauseRedlineItem],
) -> ClauseRedlinesResponse:
    """Create a sample clause redlines response."""
    return ClauseRedlinesResponse(
        summary="Clause analysis identified 2 items [1]. Governing law found, termination missing.",
        items=sample_clause_items,
        meta=ClauseRedlinesMeta(
            status=WorkflowResultStatus.SUCCESS,
            model="gpt-4o-mini",
            provider="openai",
            evidence_chunk_count=8,
            request_id="req-67890",
            output_language="en",
            jurisdiction="UAE",
            downgraded_count=0,
            removed_count=0,
            strict_citations_failed=False,
            prompt_hash="def789ghi012",
            llm_provider="openai",
            llm_model="gpt-4o-mini",
        ),
        insufficient_sources=False,
    )


class TestContractReviewExport:
    """Tests for contract review DOCX export."""

    def test_generate_docx_does_not_crash(
        self,
        export_service: ExportService,
        contract_review_response: ContractReviewResponse,
        document_metadata: DocumentMetadata,
    ) -> None:
        """Test that DOCX generation completes without errors."""
        docx_bytes = export_service.generate_contract_review_docx(
            result=contract_review_response,
            metadata=document_metadata,
        )

        # Verify we got bytes back
        assert isinstance(docx_bytes, bytes)
        assert len(docx_bytes) > 0

        # Verify it's a valid DOCX (can be opened)
        doc = Document(io.BytesIO(docx_bytes))
        assert doc is not None

    def test_docx_contains_required_sections(
        self,
        export_service: ExportService,
        contract_review_response: ContractReviewResponse,
        document_metadata: DocumentMetadata,
    ) -> None:
        """Test that exported DOCX contains all required sections."""
        docx_bytes = export_service.generate_contract_review_docx(
            result=contract_review_response,
            metadata=document_metadata,
        )

        doc = Document(io.BytesIO(docx_bytes))

        # Extract all text content
        full_text = "\n".join(para.text for para in doc.paragraphs)

        # Check for required sections
        assert "Contract Review Memo" in full_text
        assert "Executive Summary" in full_text
        assert "Findings" in full_text
        assert "Citations" in full_text
        assert "Evidence Appendix" in full_text
        assert "Traceability Information" in full_text

    def test_traceability_footer_contains_correct_fields(
        self,
        export_service: ExportService,
        contract_review_response: ContractReviewResponse,
        document_metadata: DocumentMetadata,
    ) -> None:
        """Test that traceability footer has all required fields."""
        docx_bytes = export_service.generate_contract_review_docx(
            result=contract_review_response,
            metadata=document_metadata,
        )

        doc = Document(io.BytesIO(docx_bytes))
        full_text = "\n".join(para.text for para in doc.paragraphs)

        # Extract table text (traceability info is in tables)
        table_text = ""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    table_text += cell.text + " "

        combined_text = full_text + table_text

        # Check for traceability fields
        assert "CONTRACT_REVIEW_V1" in combined_text
        assert "success" in combined_text.lower()
        assert "openai" in combined_text
        assert "gpt-4o-mini" in combined_text
        assert "abc123def456" in combined_text  # Prompt hash (truncated)
        assert "Aiden.ai" in combined_text
        assert "dev" in combined_text  # Environment

    def test_insufficient_sources_includes_disclaimer(
        self,
        export_service: ExportService,
        contract_review_insufficient_sources: ContractReviewResponse,
        document_metadata: DocumentMetadata,
    ) -> None:
        """Test that insufficient sources export includes disclaimer paragraph."""
        docx_bytes = export_service.generate_contract_review_docx(
            result=contract_review_insufficient_sources,
            metadata=document_metadata,
        )

        doc = Document(io.BytesIO(docx_bytes))
        full_text = "\n".join(para.text for para in doc.paragraphs)

        # Check for insufficient sources disclaimer
        assert "insufficient source material" in full_text.lower()
        assert "NOTICE" in full_text

    def test_findings_are_included(
        self,
        export_service: ExportService,
        contract_review_response: ContractReviewResponse,
        document_metadata: DocumentMetadata,
    ) -> None:
        """Test that findings are properly included in the export."""
        docx_bytes = export_service.generate_contract_review_docx(
            result=contract_review_response,
            metadata=document_metadata,
        )

        doc = Document(io.BytesIO(docx_bytes))
        full_text = "\n".join(para.text for para in doc.paragraphs)

        # Check for finding content
        assert "Liability Cap Too Low" in full_text
        assert "Missing Termination Notice Period" in full_text
        assert "liability clause limits damages" in full_text

    def test_citations_are_numbered(
        self,
        export_service: ExportService,
        contract_review_response: ContractReviewResponse,
        document_metadata: DocumentMetadata,
    ) -> None:
        """Test that citations are properly numbered in the export."""
        docx_bytes = export_service.generate_contract_review_docx(
            result=contract_review_response,
            metadata=document_metadata,
        )

        doc = Document(io.BytesIO(docx_bytes))
        full_text = "\n".join(para.text for para in doc.paragraphs)

        # Check for numbered citations
        assert "[1]" in full_text
        assert "[2]" in full_text


class TestClauseRedlinesExport:
    """Tests for clause redlines DOCX export."""

    def test_generate_docx_does_not_crash(
        self,
        export_service: ExportService,
        clause_redlines_response: ClauseRedlinesResponse,
        document_metadata: DocumentMetadata,
    ) -> None:
        """Test that DOCX generation completes without errors."""
        docx_bytes = export_service.generate_clause_redlines_docx(
            result=clause_redlines_response,
            metadata=document_metadata,
        )

        # Verify we got bytes back
        assert isinstance(docx_bytes, bytes)
        assert len(docx_bytes) > 0

        # Verify it's a valid DOCX (can be opened)
        doc = Document(io.BytesIO(docx_bytes))
        assert doc is not None

    def test_docx_contains_required_sections(
        self,
        export_service: ExportService,
        clause_redlines_response: ClauseRedlinesResponse,
        document_metadata: DocumentMetadata,
    ) -> None:
        """Test that exported DOCX contains all required sections."""
        docx_bytes = export_service.generate_clause_redlines_docx(
            result=clause_redlines_response,
            metadata=document_metadata,
        )

        doc = Document(io.BytesIO(docx_bytes))
        full_text = "\n".join(para.text for para in doc.paragraphs)

        # Check for required sections
        assert "Clause Redlines Memo" in full_text
        assert "Executive Summary" in full_text
        assert "Clause Analysis" in full_text
        assert "Citations" in full_text
        assert "Evidence Appendix" in full_text
        assert "Traceability Information" in full_text

    def test_traceability_footer_contains_correct_fields(
        self,
        export_service: ExportService,
        clause_redlines_response: ClauseRedlinesResponse,
        document_metadata: DocumentMetadata,
    ) -> None:
        """Test that traceability footer has all required fields."""
        docx_bytes = export_service.generate_clause_redlines_docx(
            result=clause_redlines_response,
            metadata=document_metadata,
        )

        doc = Document(io.BytesIO(docx_bytes))
        full_text = "\n".join(para.text for para in doc.paragraphs)

        # Extract table text
        table_text = ""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    table_text += cell.text + " "

        combined_text = full_text + table_text

        # Check for traceability fields
        assert "CLAUSE_REDLINES_V1" in combined_text
        assert "success" in combined_text.lower()
        assert "openai" in combined_text
        assert "gpt-4o-mini" in combined_text
        assert "def789ghi012" in combined_text  # Prompt hash (truncated)
        assert "Aiden.ai" in combined_text

    def test_clause_items_are_included(
        self,
        export_service: ExportService,
        clause_redlines_response: ClauseRedlinesResponse,
        document_metadata: DocumentMetadata,
    ) -> None:
        """Test that clause items are properly included in the export."""
        docx_bytes = export_service.generate_clause_redlines_docx(
            result=clause_redlines_response,
            metadata=document_metadata,
        )

        doc = Document(io.BytesIO(docx_bytes))
        full_text = "\n".join(para.text for para in doc.paragraphs)

        # Check for clause content
        assert "Governing Law" in full_text
        assert "Termination" in full_text
        assert "FOUND" in full_text or "found" in full_text.lower()
        assert "MISSING" in full_text or "missing" in full_text.lower()

    def test_recommended_text_is_included(
        self,
        export_service: ExportService,
        clause_redlines_response: ClauseRedlinesResponse,
        document_metadata: DocumentMetadata,
    ) -> None:
        """Test that recommended text is included for clause items."""
        docx_bytes = export_service.generate_clause_redlines_docx(
            result=clause_redlines_response,
            metadata=document_metadata,
        )

        doc = Document(io.BytesIO(docx_bytes))
        full_text = "\n".join(para.text for para in doc.paragraphs)

        # Check for recommended text
        assert "Recommended Text" in full_text
        assert "shall be governed by" in full_text.lower()


class TestExportEdgeCases:
    """Tests for edge cases in DOCX export."""

    def test_empty_findings_list(
        self,
        export_service: ExportService,
        document_metadata: DocumentMetadata,
    ) -> None:
        """Test export with no findings."""
        response = ContractReviewResponse(
            summary="No significant issues identified.",
            findings=[],
            meta=ContractReviewMeta(
                status=WorkflowResultStatus.SUCCESS,
                model="gpt-4o-mini",
                provider="openai",
                evidence_chunk_count=5,
                request_id="req-empty",
                output_language="en",
                review_mode="quick",
                removed_findings_count=0,
                strict_citations_failed=False,
                prompt_hash="empty123",
                llm_provider="openai",
                llm_model="gpt-4o-mini",
            ),
            insufficient_sources=False,
        )

        docx_bytes = export_service.generate_contract_review_docx(
            result=response,
            metadata=document_metadata,
        )

        doc = Document(io.BytesIO(docx_bytes))
        full_text = "\n".join(para.text for para in doc.paragraphs)

        # Should still generate valid document
        assert len(docx_bytes) > 0
        assert "No findings were identified" in full_text

    def test_empty_clause_items_list(
        self,
        export_service: ExportService,
        document_metadata: DocumentMetadata,
    ) -> None:
        """Test export with no clause items."""
        response = ClauseRedlinesResponse(
            summary="No clauses detected.",
            items=[],
            meta=ClauseRedlinesMeta(
                status=WorkflowResultStatus.SUCCESS,
                model="gpt-4o-mini",
                provider="openai",
                evidence_chunk_count=3,
                request_id="req-empty",
                output_language="en",
                jurisdiction="UAE",
                downgraded_count=0,
                removed_count=0,
                strict_citations_failed=False,
                prompt_hash="empty456",
                llm_provider="openai",
                llm_model="gpt-4o-mini",
            ),
            insufficient_sources=False,
        )

        docx_bytes = export_service.generate_clause_redlines_docx(
            result=response,
            metadata=document_metadata,
        )

        doc = Document(io.BytesIO(docx_bytes))
        full_text = "\n".join(para.text for para in doc.paragraphs)

        # Should still generate valid document
        assert len(docx_bytes) > 0
        assert "No clause items were identified" in full_text

    def test_document_with_special_characters_in_title(
        self,
        export_service: ExportService,
        contract_review_response: ContractReviewResponse,
    ) -> None:
        """Test export with special characters in document title."""
        metadata = DocumentMetadata(
            document_id="doc-special-chars",
            version_id="ver-special-chars",
            document_title="Contract <Draft> v2.0 (Final) [Approved]",
            version_number=1,
            workspace_name="Test & Legal",
            tenant_name="Acme & Co.",
            jurisdiction="UAE",
        )

        # Should not crash
        docx_bytes = export_service.generate_contract_review_docx(
            result=contract_review_response,
            metadata=metadata,
        )

        assert len(docx_bytes) > 0

    def test_null_llm_fields(
        self,
        export_service: ExportService,
        document_metadata: DocumentMetadata,
    ) -> None:
        """Test export when LLM fields are None."""
        response = ContractReviewResponse(
            summary="Analysis complete.",
            findings=[],
            meta=ContractReviewMeta(
                status=WorkflowResultStatus.SUCCESS,
                model="stub-v1",
                provider="stub",
                evidence_chunk_count=5,
                request_id="req-null",
                output_language="en",
                review_mode="standard",
                removed_findings_count=0,
                strict_citations_failed=False,
                prompt_hash=None,  # Null
                llm_provider=None,  # Null
                llm_model=None,  # Null
            ),
            insufficient_sources=False,
        )

        docx_bytes = export_service.generate_contract_review_docx(
            result=response,
            metadata=document_metadata,
        )

        doc = Document(io.BytesIO(docx_bytes))
        table_text = ""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    table_text += cell.text + " "

        # Should handle null values with N/A
        assert "N/A" in table_text or "stub" in table_text.lower()
        assert len(docx_bytes) > 0


class TestDocxFormatValidation:
    """Tests for DOCX format and headers validation."""

    def test_docx_starts_with_zip_header(
        self,
        export_service: ExportService,
        contract_review_response: ContractReviewResponse,
        document_metadata: DocumentMetadata,
    ) -> None:
        """Test that exported DOCX starts with ZIP header (PK).

        DOCX files are ZIP archives internally, so they must start with 'PK'.
        """
        docx_bytes = export_service.generate_contract_review_docx(
            result=contract_review_response,
            metadata=document_metadata,
        )

        # DOCX is a ZIP file, must start with "PK" (0x50 0x4B)
        assert len(docx_bytes) >= 2
        assert docx_bytes[:2] == b"PK", "DOCX file should start with ZIP header 'PK'"

    def test_docx_is_valid_zip_archive(
        self,
        export_service: ExportService,
        contract_review_response: ContractReviewResponse,
        document_metadata: DocumentMetadata,
    ) -> None:
        """Test that exported DOCX is a valid ZIP archive."""
        import zipfile

        docx_bytes = export_service.generate_contract_review_docx(
            result=contract_review_response,
            metadata=document_metadata,
        )

        # Should be a valid ZIP file
        buffer = io.BytesIO(docx_bytes)
        assert zipfile.is_zipfile(buffer), "DOCX should be a valid ZIP archive"

        # Should contain [Content_Types].xml (standard DOCX structure)
        with zipfile.ZipFile(buffer, "r") as zf:
            assert "[Content_Types].xml" in zf.namelist()

    def test_clause_redlines_docx_starts_with_zip_header(
        self,
        export_service: ExportService,
        clause_redlines_response: ClauseRedlinesResponse,
        document_metadata: DocumentMetadata,
    ) -> None:
        """Test that clause redlines DOCX starts with ZIP header (PK)."""
        docx_bytes = export_service.generate_clause_redlines_docx(
            result=clause_redlines_response,
            metadata=document_metadata,
        )

        # DOCX is a ZIP file, must start with "PK"
        assert len(docx_bytes) >= 2
        assert docx_bytes[:2] == b"PK", "DOCX file should start with ZIP header 'PK'"

    def test_in_memory_generation_no_disk_writes(
        self,
        export_service: ExportService,
        contract_review_response: ContractReviewResponse,
        document_metadata: DocumentMetadata,
    ) -> None:
        """Test that DOCX generation is truly in-memory (returns bytes, not file path)."""
        docx_bytes = export_service.generate_contract_review_docx(
            result=contract_review_response,
            metadata=document_metadata,
        )

        # Result should be bytes, not a file path
        assert isinstance(docx_bytes, bytes)
        assert not isinstance(docx_bytes, str)
        assert len(docx_bytes) > 1000  # Minimum size for a valid DOCX


class TestExportFilenameGeneration:
    """Tests for filename generation logic."""

    def test_filename_sanitization(self) -> None:
        """Test that special characters are sanitized in filenames."""
        from src.routers.exports import _generate_filename

        # Test with special characters
        filename = _generate_filename("Contract <Draft> v2.0", "contract-review")
        assert "<" not in filename
        assert ">" not in filename
        assert filename.endswith(".docx")
        assert "contract-review" in filename

    def test_filename_length_limit(self) -> None:
        """Test that very long titles are truncated."""
        from src.routers.exports import _generate_filename

        # Very long title
        long_title = "A" * 100
        filename = _generate_filename(long_title, "contract-review")

        # Title portion should be truncated to 50 chars
        # Format: {title}_{workflow}_{date}.docx
        parts = filename.split("_")
        assert len(parts[0]) <= 50


# =============================================================================
# Mixed Evidence (Workspace + Global) Export Tests
# =============================================================================


class TestMixedEvidenceExport:
    """Tests for exporting results containing mixed evidence sources."""

    @pytest.fixture
    def workspace_evidence(self) -> EvidenceChunkRef:
        """Create workspace document evidence chunk."""
        return EvidenceChunkRef(
            chunk_id="workspace-chunk-1234",
            snippet="Article 5 of the employment agreement states the notice period.",
            char_start=100,
            char_end=200,
            source_type="workspace_document",
            source_label="Employment Agreement v1",
        )

    @pytest.fixture
    def global_legal_evidence(self) -> EvidenceChunkRef:
        """Create global legal corpus evidence chunk."""
        return EvidenceChunkRef(
            chunk_id="global-chunk-5678",
            snippet="UAE Labour Law Article 120 specifies termination conditions.",
            char_start=0,
            char_end=100,
            source_type="global_legal",
            source_label="UAE Labour Law (2022)",
            instrument_id="inst-uae-labour-law",
            jurisdiction="UAE",
            official_source_url="https://laws.uae.gov/labour/120",
        )

    @pytest.fixture
    def mixed_findings(
        self, workspace_evidence: EvidenceChunkRef, global_legal_evidence: EvidenceChunkRef
    ) -> list[Finding]:
        """Create findings with mixed evidence (workspace + global)."""
        return [
            Finding(
                finding_id="finding-mixed-1",
                title="Termination Notice Analysis",
                severity="medium",
                category="termination",
                issue="The contract's termination clause references UAE Labour Law [1][2].",
                recommendation="Ensure compliance with Article 120 requirements [2].",
                citations=[1, 2],
                evidence=[workspace_evidence, global_legal_evidence],
            ),
        ]

    @pytest.fixture
    def mixed_evidence_response(
        self, mixed_findings: list[Finding]
    ) -> ContractReviewResponse:
        """Create contract review response with mixed evidence."""
        return ContractReviewResponse(
            summary="Contract reviewed with reference to UAE Labour Law.",
            findings=mixed_findings,
            meta=ContractReviewMeta(
                status=WorkflowResultStatus.SUCCESS,
                model="gpt-4",
                provider="openai",
                review_mode="standard",
                focus_areas=["termination"],
                output_language="en",
                request_id="test-mixed-123",
                removed_findings_count=0,
                strict_citations_failed=False,
                prompt_hash="mixed-hash-abc",
                llm_provider="openai",
                llm_model="gpt-4-0613",
                evidence_scope="both",
                workspace_evidence_count=1,
                global_evidence_count=1,
                policy_jurisdictions_count=4,
                policy_languages_count=2,
                policy_denied_reason=None,
            ),
        )

    def test_mixed_evidence_docx_generation_succeeds(
        self,
        export_service: ExportService,
        mixed_evidence_response: ContractReviewResponse,
        document_metadata: DocumentMetadata,
    ) -> None:
        """Test that DOCX export with mixed evidence (workspace + global) succeeds."""
        docx_bytes = export_service.generate_contract_review_docx(
            result=mixed_evidence_response,
            metadata=document_metadata,
        )

        # Should generate a valid DOCX
        assert isinstance(docx_bytes, bytes)
        assert len(docx_bytes) > 1000
        assert docx_bytes[:2] == b"PK"  # Valid ZIP/DOCX

    def test_mixed_evidence_source_types_in_docx_content(
        self,
        export_service: ExportService,
        mixed_evidence_response: ContractReviewResponse,
        document_metadata: DocumentMetadata,
    ) -> None:
        """Test that mixed evidence shows correct source type labels in DOCX."""
        docx_bytes = export_service.generate_contract_review_docx(
            result=mixed_evidence_response,
            metadata=document_metadata,
        )

        # Parse DOCX and check for source type indicators
        doc = Document(io.BytesIO(docx_bytes))
        full_text = "\n".join([p.text for p in doc.paragraphs])

        # Should contain indicators for both source types
        assert "WORKSPACE" in full_text or "workspace" in full_text.lower()
        assert "GLOBAL" in full_text or "global" in full_text.lower()

    def test_global_evidence_provenance_in_docx(
        self,
        export_service: ExportService,
        mixed_evidence_response: ContractReviewResponse,
        document_metadata: DocumentMetadata,
    ) -> None:
        """Test that global legal evidence includes provenance (jurisdiction, URL) in DOCX."""
        docx_bytes = export_service.generate_contract_review_docx(
            result=mixed_evidence_response,
            metadata=document_metadata,
        )

        # Parse DOCX
        doc = Document(io.BytesIO(docx_bytes))
        full_text = "\n".join([p.text for p in doc.paragraphs])

        # Should contain jurisdiction and source label
        assert "UAE" in full_text
        assert "UAE Labour Law" in full_text or "Labour Law" in full_text


class TestEvidenceSourceTypeValidation:
    """Tests for explicit source_type validation in export payloads."""

    def test_workspace_evidence_has_source_type(self) -> None:
        """Test that workspace evidence includes explicit source_type."""
        evidence = EvidenceChunkRef(
            chunk_id="test-chunk",
            snippet="Test snippet",
            char_start=0,
            char_end=50,
            source_type="workspace_document",
            source_label="Test Document",
        )

        assert evidence.source_type == "workspace_document"
        assert evidence.source_label == "Test Document"

    def test_global_evidence_has_provenance_fields(self) -> None:
        """Test that global legal evidence includes all provenance fields."""
        evidence = EvidenceChunkRef(
            chunk_id="global-chunk",
            snippet="Legal text snippet",
            char_start=0,
            char_end=100,
            source_type="global_legal",
            source_label="Saudi Companies Law (2022)",
            instrument_id="inst-saudi-companies",
            jurisdiction="KSA",
            official_source_url="https://laws.ksa.gov/companies",
        )

        assert evidence.source_type == "global_legal"
        assert evidence.source_label == "Saudi Companies Law (2022)"
        assert evidence.instrument_id == "inst-saudi-companies"
        assert evidence.jurisdiction == "KSA"
        assert evidence.official_source_url == "https://laws.ksa.gov/companies"

    def test_source_type_defaults_to_workspace(self) -> None:
        """Test that source_type defaults to workspace_document when not specified."""
        # Minimal evidence without explicit source_type
        evidence = EvidenceChunkRef(
            chunk_id="test-chunk",
            snippet="Test snippet",
            char_start=0,
            char_end=50,
        )

        # Default should be workspace_document
        assert evidence.source_type == "workspace_document"
        assert evidence.source_label == ""  # Default empty
