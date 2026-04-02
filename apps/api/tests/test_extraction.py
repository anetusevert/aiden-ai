"""Tests for Text Extraction and Chunking.

Test Categories:
- Integration tests: Require PostgreSQL and MinIO (with minio-init) running

Prerequisites:
    docker compose up postgres minio minio-init -d

Run all tests:
    uv run pytest tests/test_extraction.py -v

Run only extraction tests:
    uv run pytest -m integration tests/test_extraction.py -v
"""

import io

import pytest
from httpx import AsyncClient
from sqlalchemy import text

from src.storage.s3 import get_storage_client


@pytest.fixture(scope="module", autouse=True)
def check_minio_bucket():
    """Fail fast if MinIO bucket doesn't exist (minio-init not run)."""
    client = get_storage_client()
    if not client.bucket_exists():
        pytest.fail(
            f"S3 bucket '{client.bucket_name}' does not exist. "
            "Did you run minio-init? Use: docker compose up postgres minio minio-init -d"
        )


def create_minimal_docx() -> bytes:
    """Create a minimal valid DOCX file for testing.

    Uses python-docx to create a simple document with test content.
    """
    from docx import Document

    doc = Document()
    doc.add_paragraph("This is a test document.")
    doc.add_paragraph("It contains multiple paragraphs for extraction testing.")
    doc.add_paragraph("Third paragraph with more content to ensure proper chunking.")

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def create_minimal_pdf() -> bytes:
    """Create a minimal valid PDF file for testing.

    Creates a simple PDF with text content using PyMuPDF if available,
    otherwise returns a minimal PDF structure.
    """
    try:
        import fitz

        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((50, 50), "This is a test PDF document.")
        page.insert_text((50, 70), "It contains text for extraction testing.")

        buffer = io.BytesIO()
        doc.save(buffer)
        doc.close()
        buffer.seek(0)
        return buffer.read()
    except ImportError:
        # Return a minimal PDF that can be parsed (may have empty text)
        return b"""%PDF-1.4
1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj
2 0 obj << /Type /Pages /Kids [3 0 R] /Count 1 >> endobj
3 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] >> endobj
xref
0 4
0000000000 65535 f 
0000000009 00000 n 
0000000058 00000 n 
0000000115 00000 n 
trailer << /Size 4 /Root 1 0 R >>
startxref
196
%%EOF"""


@pytest.fixture
async def clean_extraction_db(clean_db):
    """Clean extraction tables before tests."""
    await clean_db.execute(text("DELETE FROM document_chunks"))
    await clean_db.execute(text("DELETE FROM document_texts"))
    await clean_db.execute(text("DELETE FROM document_versions"))
    await clean_db.execute(text("DELETE FROM documents"))
    await clean_db.commit()
    yield clean_db


async def bootstrap_and_login(async_client: AsyncClient, admin_email: str):
    """Helper to bootstrap tenant/workspace and get admin token."""
    bootstrap_response = await async_client.post(
        "/tenants",
        json={
            "name": "Extraction Test Tenant",
            "primary_jurisdiction": "UAE",
            "data_residency_policy": "UAE",
            "bootstrap": {
                "admin_user": {"password": "Testpass123", "email": admin_email, "full_name": "Extraction Admin"},
                "workspace": {"name": "Extraction Workspace"},
            },
        },
    )
    assert bootstrap_response.status_code == 201
    data = bootstrap_response.json()

    login_response = await async_client.post(
        "/auth/dev-login",
        json={
            "tenant_id": data["tenant_id"],
            "workspace_id": data["workspace_id"],
            "email": admin_email,
        },
    )
    assert login_response.status_code == 200
    token = login_response.cookies.get("access_token")
    assert token, "Expected access_token cookie from dev-login"

    return data, token


@pytest.mark.integration
class TestDocxExtraction:
    """Tests for DOCX text extraction."""

    @pytest.mark.asyncio
    async def test_upload_docx_creates_text_and_chunks(
        self,
        async_client: AsyncClient,
        clean_extraction_db,
    ):
        """Upload a DOCX file creates document_texts and document_chunks records."""
        data, token = await bootstrap_and_login(async_client, "admin@docx1.com")

        # Create a minimal DOCX file
        docx_content = create_minimal_docx()

        # Upload document
        files = {
            "file": (
                "test.docx",
                docx_content,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
        }
        form_data = {
            "title": "Test DOCX Document",
            "document_type": "contract",
            "jurisdiction": "UAE",
            "language": "en",
            "confidentiality": "internal",
        }

        response = await async_client.post(
            "/documents",
            headers={"Authorization": f"Bearer {token}"},
            files=files,
            data=form_data,
        )

        assert response.status_code == 201
        result = response.json()
        doc_id = result["document"]["id"]
        version_id = result["version"]["id"]

        # Get extracted text metadata
        text_response = await async_client.get(
            f"/documents/{doc_id}/versions/{version_id}/text",
            headers={"Authorization": f"Bearer {token}"},
        )

        assert text_response.status_code == 200
        text_result = text_response.json()

        assert text_result["version_id"] == version_id
        assert text_result["extraction_method"] == "docx"
        assert text_result["text_length"] > 0
        assert text_result["extracted_text"] is None  # Not included by default

    @pytest.mark.asyncio
    async def test_get_text_with_include_text_flag(
        self,
        async_client: AsyncClient,
        clean_extraction_db,
    ):
        """Get extracted text with include_text=true returns full text."""
        data, token = await bootstrap_and_login(async_client, "admin@docx2.com")

        docx_content = create_minimal_docx()

        files = {
            "file": (
                "test.docx",
                docx_content,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
        }
        form_data = {
            "title": "Test DOCX Include Text",
            "document_type": "memo",
            "jurisdiction": "UAE",
            "language": "en",
            "confidentiality": "public",
        }

        response = await async_client.post(
            "/documents",
            headers={"Authorization": f"Bearer {token}"},
            files=files,
            data=form_data,
        )
        result = response.json()
        doc_id = result["document"]["id"]
        version_id = result["version"]["id"]

        # Get with include_text=true
        text_response = await async_client.get(
            f"/documents/{doc_id}/versions/{version_id}/text",
            headers={"Authorization": f"Bearer {token}"},
            params={"include_text": "true"},
        )

        assert text_response.status_code == 200
        text_result = text_response.json()

        assert text_result["extracted_text"] is not None
        assert "test document" in text_result["extracted_text"].lower()

    @pytest.mark.asyncio
    async def test_get_chunks_returns_chunks(
        self,
        async_client: AsyncClient,
        clean_extraction_db,
    ):
        """Get chunks returns document chunks with correct structure."""
        data, token = await bootstrap_and_login(async_client, "admin@docx3.com")

        docx_content = create_minimal_docx()

        files = {
            "file": (
                "test.docx",
                docx_content,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
        }
        form_data = {
            "title": "Test DOCX Chunks",
            "document_type": "policy",
            "jurisdiction": "DIFC",
            "language": "en",
            "confidentiality": "confidential",
        }

        response = await async_client.post(
            "/documents",
            headers={"Authorization": f"Bearer {token}"},
            files=files,
            data=form_data,
        )
        result = response.json()
        doc_id = result["document"]["id"]
        version_id = result["version"]["id"]

        # Get chunks
        chunks_response = await async_client.get(
            f"/documents/{doc_id}/versions/{version_id}/chunks",
            headers={"Authorization": f"Bearer {token}"},
        )

        assert chunks_response.status_code == 200
        chunks_result = chunks_response.json()

        assert chunks_result["version_id"] == version_id
        assert chunks_result["document_id"] == doc_id
        assert chunks_result["chunk_count"] >= 1
        assert len(chunks_result["chunks"]) >= 1

        # Verify chunk structure
        first_chunk = chunks_result["chunks"][0]
        assert first_chunk["chunk_index"] == 0
        assert "text" in first_chunk
        assert first_chunk["char_start"] >= 0
        assert first_chunk["char_end"] > first_chunk["char_start"]


@pytest.mark.integration
class TestPdfExtraction:
    """Tests for PDF text extraction."""

    @pytest.mark.asyncio
    async def test_upload_pdf_creates_text_record(
        self,
        async_client: AsyncClient,
        clean_extraction_db,
    ):
        """Upload a PDF creates document_texts record (even if empty)."""
        data, token = await bootstrap_and_login(async_client, "admin@pdf1.com")

        pdf_content = create_minimal_pdf()

        files = {
            "file": ("test.pdf", pdf_content, "application/pdf"),
        }
        form_data = {
            "title": "Test PDF Document",
            "document_type": "regulatory",
            "jurisdiction": "UAE",
            "language": "en",
            "confidentiality": "internal",
        }

        response = await async_client.post(
            "/documents",
            headers={"Authorization": f"Bearer {token}"},
            files=files,
            data=form_data,
        )

        assert response.status_code == 201
        result = response.json()
        doc_id = result["document"]["id"]
        version_id = result["version"]["id"]

        # Get extracted text metadata
        text_response = await async_client.get(
            f"/documents/{doc_id}/versions/{version_id}/text",
            headers={"Authorization": f"Bearer {token}"},
        )

        # Should have text record even if extraction produced empty text
        assert text_response.status_code == 200
        text_result = text_response.json()
        assert text_result["version_id"] == version_id
        assert text_result["extraction_method"] in ["pymupdf", "pdfminer", "unsupported"]


@pytest.mark.integration
class TestExtractionAuditLogs:
    """Tests for extraction audit logging."""

    @pytest.mark.asyncio
    async def test_successful_extraction_creates_audit_log(
        self,
        async_client: AsyncClient,
        clean_extraction_db,
    ):
        """Successful extraction creates document.extract.success audit log."""
        data, token = await bootstrap_and_login(async_client, "admin@audit1.com")

        docx_content = create_minimal_docx()

        files = {
            "file": (
                "audit_test.docx",
                docx_content,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
        }
        form_data = {
            "title": "Audit Test DOCX",
            "document_type": "other",
            "jurisdiction": "UAE",
            "language": "en",
            "confidentiality": "public",
        }

        response = await async_client.post(
            "/documents",
            headers={"Authorization": f"Bearer {token}"},
            files=files,
            data=form_data,
        )
        assert response.status_code == 201

        # Check audit log for extraction success
        audit_response = await async_client.get(
            "/audit-logs",
            headers={"Authorization": f"Bearer {token}"},
            params={"action": "document.extract.success"},
        )

        assert audit_response.status_code == 200
        audit_data = audit_response.json()

        extract_event = next(
            (e for e in audit_data["items"] if e["action"] == "document.extract.success"),
            None,
        )
        assert extract_event is not None
        assert extract_event["status"] == "success"
        assert "method" in extract_event.get("meta", {})
        assert "chunk_count" in extract_event.get("meta", {})


@pytest.mark.integration
class TestExtractionTenantIsolation:
    """Tests for tenant/workspace isolation in extraction."""

    @pytest.mark.asyncio
    async def test_cannot_access_other_workspace_text(
        self,
        async_client: AsyncClient,
        clean_extraction_db,
    ):
        """Cannot access extracted text from another workspace."""
        # Create first workspace
        data1, token1 = await bootstrap_and_login(async_client, "admin@ws1ext.com")

        docx_content = create_minimal_docx()

        files = {
            "file": (
                "ws1.docx",
                docx_content,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
        }
        form_data = {
            "title": "WS1 Document",
            "document_type": "contract",
            "jurisdiction": "UAE",
            "language": "en",
            "confidentiality": "internal",
        }

        response1 = await async_client.post(
            "/documents",
            headers={"Authorization": f"Bearer {token1}"},
            files=files,
            data=form_data,
        )
        result1 = response1.json()
        doc_id = result1["document"]["id"]
        version_id = result1["version"]["id"]

        # Create second workspace
        data2, token2 = await bootstrap_and_login(async_client, "admin@ws2ext.com")

        # Try to access text from other workspace
        text_response = await async_client.get(
            f"/documents/{doc_id}/versions/{version_id}/text",
            headers={"Authorization": f"Bearer {token2}"},
        )

        assert text_response.status_code == 404

        # Try to access chunks from other workspace
        chunks_response = await async_client.get(
            f"/documents/{doc_id}/versions/{version_id}/chunks",
            headers={"Authorization": f"Bearer {token2}"},
        )

        assert chunks_response.status_code == 404


@pytest.mark.integration
class TestChunkerUnit:
    """Unit tests for the chunker module."""

    def test_empty_text_returns_no_chunks(self):
        """Empty text produces no chunks."""
        from src.extraction import create_chunks

        chunks = create_chunks("")
        assert chunks == []

        chunks = create_chunks("   \n\n   ")
        assert chunks == []

    def test_short_text_produces_single_chunk(self):
        """Short text produces a single chunk."""
        from src.extraction import create_chunks

        text = "This is a short piece of text."
        chunks = create_chunks(text)

        assert len(chunks) == 1
        assert chunks[0].text == text
        assert chunks[0].chunk_index == 0
        assert chunks[0].char_start == 0
        assert chunks[0].char_end == len(text)

    def test_long_text_produces_multiple_chunks(self):
        """Long text produces multiple chunks."""
        from src.extraction import create_chunks

        # Create text longer than max chunk size
        text = ("This is paragraph one. " * 50 + "\n\n") * 5
        chunks = create_chunks(text)

        assert len(chunks) > 1
        # Verify indices are sequential
        for i, chunk in enumerate(chunks):
            assert chunk.chunk_index == i

    def test_chunks_are_deterministic(self):
        """Same input produces identical chunks."""
        from src.extraction import create_chunks

        text = "First paragraph.\n\nSecond paragraph.\n\nThird paragraph."
        chunks1 = create_chunks(text)
        chunks2 = create_chunks(text)

        assert len(chunks1) == len(chunks2)
        for c1, c2 in zip(chunks1, chunks2):
            assert c1.text == c2.text
            assert c1.char_start == c2.char_start
            assert c1.char_end == c2.char_end

    def test_arabic_text_preserved(self):
        """Arabic text is preserved without modification."""
        from src.extraction import create_chunks

        arabic_text = "هذا نص عربي للاختبار.\n\nهذه فقرة ثانية."
        chunks = create_chunks(arabic_text)

        assert len(chunks) >= 1
        # Verify text is preserved
        reconstructed = arabic_text[chunks[0].char_start : chunks[0].char_end]
        assert chunks[0].text == reconstructed
